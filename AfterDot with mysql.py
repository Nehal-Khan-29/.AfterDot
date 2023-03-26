'''
project done by 
NEHAL KHAN
SRIDHARAN
ATHISH VISHNU
'''

import tkinter as tk
from tkinter import *
from tkinter import messagebox
from tkinter import filedialog as fd
from tkinter.filedialog import askdirectory
from tkinter import ttk as ttk
from datetime import date
import time
import os
from PIL import ImageTk,Image
from PIL import Image
from pdf2docx import Converter
from pdf2docx import parse
from typing import Tuple
from docx2pdf import convert
import aspose.words as aw
import docx
from docx import Document
from PyPDF2 import PdfMerger, PdfReader
from pdf2image import convert_from_path
import PyPDF2
import mysql.connector


mydb=mysql.connector.connect(host='localhost',user='root',password='password',database='afterdot')


# Page Close Confirmations (Messagebox):

def homelogout():
    messagebox.showinfo('Thank You','Logged out successfully')
    home.destroy()
        
def homeclose():
    if messagebox.askokcancel('Quit','Do you want to logout and quit?'):
        home.destroy()
        quit()

# icon window

window = tk.Tk()
image = Image.open("E:\\NK programs\\Python\\python save\\AfterDot\\icon.png")
tk_image = ImageTk.PhotoImage(image)
image_label = tk.Label(window, image=tk_image)
image_label.pack()
window.update()
screen_width = window.winfo_screenwidth()
screen_height = window.winfo_screenheight()
window_width = window.winfo_width()
window_height = window.winfo_height()
x = int((screen_width - window_width) / 2)
y = int((screen_height - window_height) / 2)
window.geometry("+{}+{}".format(x, y))
window.after(2000)
window.destroy()


#pdf to text:

def pdftotxt():
    root = tk.Tk()
    root.withdraw()

    filetypes = [('PDF files', '*.pdf'), ('All files', '*.*')]
    filepaths = fd.askopenfilenames(filetypes=filetypes)

    if len(filepaths) > 0:
        filepath = filepaths[0]

        with open(filepath, 'rb') as pdf_file:

            pdf_reader = PyPDF2.PdfReader(pdf_file)
            num_pages = len(pdf_reader.pages)
            text = ""

            for i in range(num_pages):
                page = pdf_reader.pages[i]
                text += page.extract_text()

        filename_without_ext = os.path.splitext(os.path.basename(filepath))[0]

        txt_filepath = os.path.join(os.path.dirname(filepath), f"{filename_without_ext}.txt")
        with open(txt_filepath, 'w', encoding='utf-8') as txt_file:
            txt_file.write(text)

        print("Successfully converted")
    else:
        print("No files selected")

    query1="select points from points where sno={}".format(8)
    mycur=mydb.cursor()
    mycur.execute(query1)
    row=mycur.fetchone()
    mydb.commit()
    query="update points set points={} where sno={}".format(row[0]+1,8)
    mycur=mydb.cursor()
    mycur.execute(query)
    mydb.commit()


# Pdf to word 2:

def pdftoword2():
    pdf_filepath = fd.askopenfilename(filetypes=(('PDF files', '*.pdf'), ('All files', '*.*')))

    with open(pdf_filepath, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        pdf_text = ''
        for page in range(len(pdf_reader.pages)):
            pdf_text += pdf_reader.pages[page].extract_text()

    doc = docx.Document()
    doc.add_paragraph(pdf_text)
    docx_filepath = pdf_filepath.replace('.pdf', '.docx')
    doc.save(docx_filepath)

    query1="select points from points where sno={}".format(9)
    mycur=mydb.cursor()
    mycur.execute(query1)
    row=mycur.fetchone()
    mydb.commit()
    query="update points set points={} where sno={}".format(row[0]+1,9)
    mycur=mydb.cursor()
    mycur.execute(query)
    mydb.commit()


# Word To PDF Page:

def wordtopdf(): 
    filetype=('docx files','*.docx'),('all files','*.*')
    filepath=fd.askopenfilenames(filetypes=filetype)
    a=filepath

    z=a[0]
    b=z[-6:-len(z)-1:-1]
    c=b[-1:-len(b)-1:-1]

    if z[-5:] == '.docx':
        convert(f"{c}.docx",
        f"{c}.pdf")

    query1="select points from points where sno={}".format(1)
    mycur=mydb.cursor()
    mycur.execute(query1)
    row=mycur.fetchone()
    mydb.commit()
    query="update points set points={} where sno={}".format(row[0]+1,1)
    mycur=mydb.cursor()
    mycur.execute(query)
    mydb.commit()


    def completed():   
        Label(home,text=('''      
            '''),font=('Arial',12),bg='Lightgreen',borderwidth=0,relief='solid').place(x=1030,rely=0.85)    
        Label(home,text=('''      
            '''),font=('Arial',12),bg='red',borderwidth=0,relief='solid').place(x=1080,rely=0.85)    
    completed()



# PDF To Word Page:

def pdftoword():
    filetype=('pdf files','*.pdf'),('all files','*.*')
    filepath=fd.askopenfilenames(filetypes=filetype)
    a=filepath

    z=a[0]
    b=z[-5:-len(z)-1:-1]
    c=b[-1:-len(b)-1:-1]

    parse(f"{c}.pdf",f"{c}.docx")

    query1="select points from points where sno={}".format(2)
    mycur=mydb.cursor()
    mycur.execute(query1)
    row=mycur.fetchone()
    mydb.commit()
    query="update points set points={} where sno={}".format(row[0]+1,2)
    mycur=mydb.cursor()
    mycur.execute(query)
    mydb.commit()

# JPG To PDF Page:

def jpgtopdf():

    filetype=('jpg files','*.jpg'),('all files','*.*')
    filepath=fd.askopenfilenames(filetypes=filetype)
    a=filepath

    z=a[0]
    b=z[-5:-len(z)-1:-1]
    c=b[-1:-len(b)-1:-1]
    
    image = Image.open(z,'r')
    im1 = image.convert('RGB')
    im1.save(f"{c}.pdf")

    query1="select points from points where sno={}".format(6)
    mycur=mydb.cursor()
    mycur.execute(query1)
    row=mycur.fetchone()
    mydb.commit()
    query="update points set points={} where sno={}".format(row[0]+1,6)
    mycur=mydb.cursor()
    mycur.execute(query)
    mydb.commit()

# PDF To JPG Page:

def pdftojpg():

    filetype=('pdf files','*.pdf'),('all files','*.*')
    filepath=fd.askopenfilenames(filetypes=filetype)
    a=filepath

    z=a[0]
    b=z[-5:-len(z)-1:-1]
    c=b[-1:-len(b)-1:-1]

    doc = aw.Document(z)
            
    for page in range(0, doc.page_count):
        extractedPage = doc.extract_pages(page, 1)
        extractedPage.save(f"{z+str(page+1)}.jpg")

    query1="select points from points where sno={}".format(5)
    mycur=mydb.cursor()
    mycur.execute(query1)
    row=mycur.fetchone()
    mydb.commit()
    query="update points set points={} where sno={}".format(row[0]+1,5)
    mycur=mydb.cursor()
    mycur.execute(query)
    mydb.commit()


# 2xPDF TO PDF

def pdftopdf():
    filetype1=('pdf files','*.pdf'),('all files','*.*')
    filepath1=fd.askopenfilenames(filetypes=filetype1)
    a=filepath1
    filetype2=('pdf files','*.pdf'),('all files','*.*')
    filepath2=fd.askopenfilenames(filetypes=filetype2)
    x=filepath2

    z=a[0]
    y=x[0]
    b=z[-5:-len(z)-1:-1]
    c=b[-1:-len(b)-1:-1]

    merger=PdfMerger()
    merger.append(z)
    merger.append(y)
    merger.write(c+"merged.pdf")
    merger.close()

    query1="select points from points where sno={}".format(7)
    mycur=mydb.cursor()
    mycur.execute(query1)
    row=mycur.fetchone()
    mydb.commit()
    query="update points set points={} where sno={}".format(row[0]+1,7)
    mycur=mydb.cursor()
    mycur.execute(query)
    mydb.commit()

# PNG TO JPG

def pngtojpg():
    filetype=('png files','*.png'),('all files','*.*')
    filepath=fd.askopenfilenames(filetypes=filetype)
    a=filepath

    z=a[0]
    b=z[-5:-len(z)-1:-1]
    c=b[-1:-len(b)-1:-1]
    
    image = Image.open(z,'r')
    im1 = image.convert('RGB')
    im1.save(f"{c}.jpg")

    query1="select points from points where sno={}".format(3)
    mycur=mydb.cursor()
    mycur.execute(query1)
    row=mycur.fetchone()
    mydb.commit()
    query="update points set points={} where sno={}".format(row[0]+1,3)
    mycur=mydb.cursor()
    mycur.execute(query)
    mydb.commit()

# JPG TO PNG

def jpgtopng():
    filetype=('jpg files','*.jpg'),('all files','*.*')
    filepath=fd.askopenfilenames(filetypes=filetype)
    a=filepath

    z=a[0]
    b=z[-5:-len(z)-1:-1]
    c=b[-1:-len(b)-1:-1]
    
    image = Image.open(z,'r')
    im1 = image.convert('RGB')
    im1.save(f"{c}.png")

    query1="select points from points where sno={}".format(4)
    mycur=mydb.cursor()
    mycur.execute(query1)
    row=mycur.fetchone()
    mydb.commit()
    query="update points set points={} where sno={}".format(row[0]+1,4)
    mycur=mydb.cursor()
    mycur.execute(query)
    mydb.commit()

# tree view:

def treeview(page,h,px,py):

    global tree
    tree=ttk.Treeview(page,column=('#c1','#c2','#c3'),show='headings',height=h)

    tree.column('#1',width=140,minwidth=140,anchor=tk.CENTER)
    tree.column('#2',width=140,minwidth=140,anchor=tk.CENTER)
    tree.column('#3',width=140,minwidth=140,anchor=tk.CENTER)

    tree.heading('#1',text='sno')
    tree.heading('#2',text='func')
    tree.heading('#3',text='points')

    tree.place(x=px,y=py)    


# Points Page:

def pointspage():

    point=tk.Toplevel()
    point.geometry('1366x768')
    point.title('AfterDot - Points')
    point.state('zoomed')

    pointpic=ImageTk.PhotoImage(Image.open("E:\\NK programs\\Python\\python save\\AfterDot\\afterdot front.png"))
    pointpanel=Label(point,image=pointpic)
    pointpanel.pack(side='top',fill='both',expand='yes')

    treeview(point,9,580,220)

    global ROWS
    mycur=mydb.cursor()
    mycur.execute('SELECT * FROM points')
    ROWS=mycur.fetchall()
    for ROW in ROWS:
        tree.insert('',tk.END,values=ROW)

    def reset():
        for i in range(1,10):
            query="update points set points={} where sno={}".format(0,i)
            mycur=mydb.cursor()
            mycur.execute(query)
            mydb.commit()
            point.destroy()   

    Button(point,text='Reset points',font=('Arial',20),command=reset,height=1,width=18,bg='red',
       fg='white',activebackground='Skyblue',activeforeground='thistle1').place(x=1030,y=580)

    point.mainloop()

# About Page:

def aboutpage():

    about=tk.Toplevel()
    about.geometry('1366x768')
    about.title('AfterDot - About')
    about.state('zoomed')

    aboutpic=ImageTk.PhotoImage(Image.open("E:\\NK programs\\Python\\python save\\AfterDot\\afterdot front.png"))
    aboutpanel=Label(about,image=aboutpic)
    aboutpanel.pack(side='top',fill='both',expand='yes')

    def help():
        Label(about,text='''MAIL ID                         : nehalmicro29@gmail.com''',font=('Arial',16)).place(x=750,y=550)

    Button(about,text='Contact Us',font=('Arial',20),command=help,height=1,width=16,bg='white',
    fg='gray6',activebackground='Skyblue',activeforeground='thistle1').place(x=750,y=500)

    Label(about,text=('''This program is created on 16 january 2023 by Nehal Khan.'''),font=('Arial',15),bg='white').place(x=250,y=200)


    about.mainloop()


# Home Page:

home=tk.Tk()
home.geometry('1366x768')
home.title('AfterDot')
home.state('zoomed')
home.protocol('WM_DELETE_WINDOW',homeclose)

currtime=time.strftime('%H:%M')
currdate=date.today().strftime("%d/%m/%Y")

homepic=ImageTk.PhotoImage(Image.open("E:\\NK programs\\Python\\python save\\AfterDot\\afterdot front.png"))
homepanel=Label(home,image=homepic)
homepanel.pack(side='top',fill='both',expand='yes')

Label(home,text=('Logged in: '+currtime+' - '+currdate),font=('Arial',16),bg='blue',fg="white").place(x=1030,y=100)

Label(home,text=(''' Convert them as many times as you like !!! '''),
                 font=('Arial',12),bg='white',borderwidth=0,relief='solid').place(relx=0.41,rely=0.85)

Button(home,text='PDF -> TEXT',font=('Arial',20),command=pdftotxt,height=1,width=16,bg='white',
       fg='gray6',activebackground='Skyblue',activeforeground='thistle1').place(x=650,y=460)
Button(home,text='Word -> PDF',font=('Arial',20),command=wordtopdf,height=1,width=16,bg='white',
       fg='gray6',activebackground='Skyblue',activeforeground='thistle1').place(x=270,y=220)
Button(home,text='JPG -> PDF',font=('Arial',20),command=jpgtopdf,height=1,width=16,bg='white',
       fg='gray6',activebackground='Skyblue',activeforeground='thistle1').place(x=270,y=340)
Button(home,text='2xPDF -> PDF',font=('Arial',20),command=pdftopdf,height=1,width=16,bg='white',
       fg='gray6',activebackground='Skyblue',activeforeground='thistle1').place(x=270,y=460)
Button(home,text='Points',font=('Arial',20),command=pointspage,height=1,width=16,bg='blue',
       fg='white',activebackground='Skyblue',activeforeground='thistle1').place(x=270,y=580)

Button(home,text='PDF -> Word',font=('Arial',20),command=pdftoword,height=1,width=16,bg='white',
       fg='gray6',activebackground='Skyblue',activeforeground='thistle1').place(x=650,y=220)
Label(home,text=("Only borderless pdf"),font=('Arial',16),bg='white',fg="blue").place(x=690,y=275)
Button(home,text='PDF -> JPG',font=('Arial',20),command=pdftojpg,height=1,width=16,bg='white',
       fg='gray6',activebackground='Skyblue',activeforeground='thistle1').place(x=650,y=340)
Button(home,text='About Us',font=('Arial',20,),command=aboutpage,height=1,width=16,bg='blue',
       fg='white',activebackground='Skyblue',activeforeground='thistle1').place(x=650,y=580)
       
Button(home,text='PNG -> JPG',font=('Arial',20),command=pngtojpg,height=1,width=16,bg='white',
       fg='gray6',activebackground='Skyblue',activeforeground='thistle1').place(x=1030,y=220)
Button(home,text='JPG -> PNG',font=('Arial',20),command=jpgtopng,height=1,width=16,bg='white',
       fg='gray6',activebackground='Skyblue',activeforeground='thistle1').place(x=1030,y=340)
Button(home,text='Logout',font=('Arial',20),command=homelogout,height=1,width=16,bg='lightgreen',
       fg='red',activebackground='Skyblue',activeforeground='thistle1').place(x=1030,y=580)

Button(home,text='PDF -> Word 2',font=('Arial',20),command=pdftoword2,height=1,width=16,bg='white',
       fg='gray6',activebackground='Skyblue',activeforeground='thistle1').place(x=1030,y=460)
Label(home,text=("Only bordered pdf"),font=('Arial',16),bg='white',fg="blue").place(x=1070,y=515)

home.mainloop()
